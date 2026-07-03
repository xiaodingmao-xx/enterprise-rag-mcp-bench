"""DOCX document loader."""

from __future__ import annotations

from pathlib import Path
from typing import List

from src.core.types import Document
from src.libs.loader.base_loader import BaseLoader

try:
    from docx import Document as DocxDocument

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class DocxLoader(BaseLoader):
    """Load modern Word .docx files."""

    file_type = "docx"

    def __init__(self) -> None:
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                "Install with: pip install python-docx"
            )

    def load(self, file_path: str | Path) -> Document:
        path = self._validate_file(file_path)
        doc = DocxDocument(str(path))

        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            table_text = self._table_to_markdown(table)
            if table_text:
                parts.append(table_text)

        text = "\n\n".join(parts).strip()
        if not text:
            raise ValueError(f"DOCX document is empty: {path}")

        title = (doc.core_properties.title or "").strip() or path.name

        return Document(
            id=self._document_id(path),
            text=text,
            metadata={
                "source_path": str(path),
                "file_name": path.name,
                "file_type": self.file_type,
                "doc_type": self.file_type,
                "title": title,
            },
        )

    @staticmethod
    def _table_to_markdown(table: object) -> str:
        rows: list[list[str]] = []
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                rows.append(values)

        if not rows:
            return ""

        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = normalized[0]
        separator = ["---"] * width
        body = normalized[1:]

        markdown_rows: List[str] = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |",
        ]
        markdown_rows.extend("| " + " | ".join(row) + " |" for row in body)
        return "\n".join(markdown_rows)
